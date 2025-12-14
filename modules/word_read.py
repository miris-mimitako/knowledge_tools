"""Wordファイル読み込みとテキスト抽出モジュール"""
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
from pydantic import BaseModel, Field


class WordParagraphData(BaseModel):
    """Word段落データのPydanticモデル"""
    paragraph_index: int = Field(..., ge=0, description="段落インデックス")
    text: str = Field(..., description="段落のテキスト")
    style: Optional[str] = Field(None, description="段落スタイル")
    alignment: Optional[str] = Field(None, description="段落の配置")


class WordTableRowData(BaseModel):
    """Wordテーブル行データのPydanticモデル"""
    row_index: int = Field(..., ge=0, description="行インデックス")
    cells: List[str] = Field(..., description="セルのテキストリスト")


class WordTableData(BaseModel):
    """WordテーブルデータのPydanticモデル"""
    table_index: int = Field(..., ge=0, description="テーブルインデックス")
    row_count: int = Field(..., ge=0, description="行数")
    column_count: int = Field(..., ge=0, description="列数")
    rows: List[WordTableRowData] = Field(..., description="行データのリスト")
    full_text: str = Field(..., description="テーブルの全テキスト")


class WordFileData(BaseModel):
    """WordファイルデータのPydanticモデル"""
    file_path: str = Field(..., description="ファイルの絶対パス")
    filename: str = Field(..., description="ファイル名")
    paragraph_count: int = Field(..., ge=0, description="段落数")
    table_count: int = Field(..., ge=0, description="テーブル数")
    paragraphs: List[WordParagraphData] = Field(..., description="段落データのリスト")
    tables: List[WordTableData] = Field(..., description="テーブルデータのリスト")
    full_text: str = Field(..., description="全テキスト（段落とテーブルを結合）")
    word_count: int = Field(..., ge=0, description="単語数")
    extracted_at: str = Field(..., description="抽出日時（ISO形式）")
    
    class Config:
        json_schema_extra = {
            "example": {
                "file_path": "/path/to/file.docx",
                "filename": "file.docx",
                "paragraph_count": 10,
                "table_count": 2,
                "paragraphs": [],
                "tables": [],
                "full_text": "Sample text...",
                "word_count": 100,
                "extracted_at": "2024-01-01T00:00:00"
            }
        }


def read_word_file(file_path: str) -> WordFileData:
    """
    Wordファイルを読み込んでテキスト情報を抽出
    
    Args:
        file_path: Wordファイルのパス
        
    Returns:
        抽出されたテキスト情報のPydanticモデル
        
    Raises:
        FileNotFoundError: ファイルが存在しない場合
        ValueError: ファイルがWord形式でない場合
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docxが必要です。'pip install python-docx'でインストールしてください。")
    
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")
    
    if not path.suffix.lower() in ['.docx', '.doc']:
        raise ValueError(f"Wordファイルではありません: {file_path}")
    
    try:
        # Wordファイルを読み込む
        doc = Document(file_path)
        
        # 段落データを抽出
        paragraphs_data = []
        paragraph_texts = []
        
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # 空でない段落のみ追加
                paragraphs_data.append(WordParagraphData(
                    paragraph_index=idx,
                    text=text,
                    style=paragraph.style.name if paragraph.style else None,
                    alignment=str(paragraph.alignment) if paragraph.alignment else None
                ))
                paragraph_texts.append(text)
        
        # テーブルデータを抽出
        tables_data = []
        table_texts = []
        
        for table_idx, table in enumerate(doc.tables):
            rows_data = []
            table_full_text = []
            
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                cells_text = " | ".join([cell for cell in cells if cell])
                
                if cells_text:  # 空行でない場合のみ追加
                    rows_data.append(WordTableRowData(
                        row_index=row_idx,
                        cells=cells
                    ))
                    table_full_text.append(cells_text)
            
            if rows_data:
                table_text = " | ".join(table_full_text)
                tables_data.append(WordTableData(
                    table_index=table_idx,
                    row_count=len(rows_data),
                    column_count=len(table.columns),
                    rows=rows_data,
                    full_text=table_text
                ))
                table_texts.append(table_text)
        
        # 全テキストを結合
        all_texts = paragraph_texts + table_texts
        full_text = "\n".join(all_texts)
        
        return WordFileData(
            file_path=str(path.absolute()),
            filename=path.name,
            paragraph_count=len(paragraphs_data),
            table_count=len(tables_data),
            paragraphs=paragraphs_data,
            tables=tables_data,
            full_text=full_text,
            word_count=len(full_text.split()) if full_text else 0,
            extracted_at=datetime.now().isoformat()
        )
    except Exception as e:
        raise ValueError(f"Wordファイルの読み込みエラー: {str(e)}")


def read_word_simple(file_path: str) -> str:
    """
    Wordファイルを読み込んで、シンプルなテキスト文字列として返す
    
    Args:
        file_path: Wordファイルのパス
        
    Returns:
        抽出されたテキスト（改行区切り）
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docxが必要です。'pip install python-docx'でインストールしてください。")
    
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")
    
    try:
        doc = Document(file_path)
        texts = []
        
        # 段落のテキスト
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                texts.append(text)
        
        # テーブルのテキスト
        for table in doc.tables:
            table_texts = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join([cell for cell in cells if cell])
                if row_text:
                    table_texts.append(row_text)
            if table_texts:
                texts.append("=== Table ===")
                texts.extend(table_texts)
        
        return "\n".join(texts)
    except Exception as e:
        raise ValueError(f"Wordファイルの読み込みエラー: {str(e)}")


def get_word_metadata(file_path: str) -> Dict[str, Any]:
    """
    Wordファイルのメタ情報を取得（ファイルを開かずに）
    
    Args:
        file_path: Wordファイルのパス
        
    Returns:
        メタ情報の辞書
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docxが必要です。'pip install python-docx'でインストールしてください。")
    
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")
    
    try:
        # ファイルを開いてメタ情報のみ取得
        doc = Document(file_path)
        
        # 段落とテーブルの数をカウント
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        
        return {
            "file_path": str(path.absolute()),
            "filename": path.name,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "has_headers": any(p.style.name.startswith('Heading') for p in doc.paragraphs),
            "has_footers": len(doc.sections[0].footer.paragraphs) > 0 if doc.sections else False,
            "has_headers_footers": len(doc.sections[0].header.paragraphs) > 0 if doc.sections else False
        }
    except Exception as e:
        raise ValueError(f"Wordファイルのメタ情報取得エラー: {str(e)}")


def extract_word_to_chunks(file_path: str, chunk_size: int = 1000) -> List[Dict[str, Any]]:
    """
    Wordファイルを読み込んで、チャンク単位でテキストを抽出
    
    Args:
        file_path: Wordファイルのパス
        chunk_size: チャンクサイズ（文字数）
        
    Returns:
        チャンクのリスト
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docxが必要です。'pip install python-docx'でインストールしてください。")
    
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")
    
    doc = Document(file_path)
    chunks = []
    chunk_id = 0
    
    current_chunk = ""
    current_chunk_paragraphs = []
    
    # 段落を処理
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        
        if len(current_chunk) + len(text) > chunk_size and current_chunk:
            # チャンクを保存
            chunks.append({
                "chunk_id": chunk_id,
                "type": "paragraph",
                "paragraph_start": current_chunk_paragraphs[0] if current_chunk_paragraphs else para_idx,
                "paragraph_end": current_chunk_paragraphs[-1] if current_chunk_paragraphs else para_idx,
                "text": current_chunk.strip(),
                "char_count": len(current_chunk)
            })
            chunk_id += 1
            current_chunk = ""
            current_chunk_paragraphs = []
        
        current_chunk += text + "\n"
        current_chunk_paragraphs.append(para_idx)
    
    # テーブルを処理
    for table_idx, table in enumerate(doc.tables):
        table_text = ""
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join([cell for cell in cells if cell])
            if row_text:
                table_text += row_text + "\n"
        
        if table_text:
            if len(current_chunk) + len(table_text) > chunk_size and current_chunk:
                # 現在のチャンクを保存
                chunks.append({
                    "chunk_id": chunk_id,
                    "type": "paragraph",
                    "paragraph_start": current_chunk_paragraphs[0] if current_chunk_paragraphs else 0,
                    "paragraph_end": current_chunk_paragraphs[-1] if current_chunk_paragraphs else 0,
                    "text": current_chunk.strip(),
                    "char_count": len(current_chunk)
                })
                chunk_id += 1
                current_chunk = ""
                current_chunk_paragraphs = []
            
            current_chunk += f"=== Table {table_idx} ===\n{table_text}"
    
    # 最後のチャンクを保存
    if current_chunk.strip():
        chunks.append({
            "chunk_id": chunk_id,
            "type": "paragraph",
            "paragraph_start": current_chunk_paragraphs[0] if current_chunk_paragraphs else 0,
            "paragraph_end": current_chunk_paragraphs[-1] if current_chunk_paragraphs else 0,
            "text": current_chunk.strip(),
            "char_count": len(current_chunk)
        })
    
    return chunks

